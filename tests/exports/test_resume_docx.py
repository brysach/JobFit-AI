# tests/exports/test_resume_docx.py

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from src.export.resume_docx import export_materials_to_docx


def sample_user_profile() -> dict:
    return {
        "name": "Bryan Estrada",
        "email": "bryan@example.com",
        "phone_number": "555-123-4567",
        "university": "University of California, Riverside",
        "degree": "B.S. Computer Science",
    }


def sample_materials_data() -> dict:
    return {
        "resume": {
            "skills": ["Python", "Git"],
            "projects": ["Built JobFit-AI using Python."],
            "experience": ["Tutored students in problem-solving."],
        },
        "cover_letter": "Dear TechStart Hiring Team,\nI am excited to apply.",
        "strengths": ["Strong Python and Git match for this role."],
        "weaknesses": ["Practice assessment problems before the interview."],
    }


def _document_text(file_path: str | Path) -> str:
    document = Document(str(file_path))
    lines = []

    for paragraph in document.paragraphs:
        if paragraph.text:
            lines.append(paragraph.text)

    return "\n".join(lines)


def test_export_materials_to_docx_creates_file(tmp_path):
    file_path = export_materials_to_docx(
        materials_data=sample_materials_data(),
        user_profile=sample_user_profile(),
        output_dir=tmp_path,
        filename="materials.docx",
    )

    assert Path(file_path).exists()


def test_export_materials_to_docx_contains_resume_cover_letter_strengths_and_weaknesses(
    tmp_path,
):
    file_path = export_materials_to_docx(
        materials_data=sample_materials_data(),
        user_profile=sample_user_profile(),
        output_dir=tmp_path,
        filename="materials.docx",
    )

    text = _document_text(file_path)

    assert "Bryan Estrada" in text
    assert "bryan@example.com | 555-123-4567" in text
    assert "EDUCATION" in text
    assert "University of California, Riverside - B.S. Computer Science" in text
    assert "SKILLS" in text
    assert "Python" in text
    assert "PROJECTS" in text
    assert "Built JobFit-AI using Python." in text
    assert "EXPERIENCE" in text
    assert "Tutored students in problem-solving." in text
    assert "COVER LETTER" in text
    assert "Dear TechStart Hiring Team," in text
    assert "STRENGTHS" in text
    assert "Strong Python and Git match for this role." in text
    assert "WEAKNESSES" in text
    assert "Practice assessment problems before the interview." in text


def test_export_materials_to_docx_applies_name_font_formatting(tmp_path):
    file_path = export_materials_to_docx(
        materials_data=sample_materials_data(),
        user_profile=sample_user_profile(),
        output_dir=tmp_path,
        filename="materials.docx",
    )

    document = Document(str(file_path))
    name_run = document.paragraphs[0].runs[0]

    assert name_run.text == "Bryan Estrada"
    assert name_run.bold is True
    assert name_run.font.name == "Garamond"
    assert name_run.font.size == Pt(28)


def test_export_materials_to_docx_applies_contact_font_formatting(tmp_path):
    file_path = export_materials_to_docx(
        materials_data=sample_materials_data(),
        user_profile=sample_user_profile(),
        output_dir=tmp_path,
        filename="materials.docx",
    )

    document = Document(str(file_path))
    contact_run = document.paragraphs[1].runs[0]

    assert contact_run.text == "bryan@example.com | 555-123-4567"
    assert contact_run.bold is None or contact_run.bold is False
    assert contact_run.font.name == "Garamond"
    assert contact_run.font.size == Pt(13)


def test_export_materials_to_docx_adds_separator_lines(tmp_path):
    file_path = export_materials_to_docx(
        materials_data=sample_materials_data(),
        user_profile=sample_user_profile(),
        output_dir=tmp_path,
        filename="materials.docx",
    )

    document = Document(str(file_path))

    separator_count = 0

    for paragraph in document.paragraphs:
        paragraph_properties = paragraph._p.pPr

        if paragraph_properties is None:
            continue

        borders = paragraph_properties.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr"
        )

        if borders is not None:
            separator_count += 1

    assert separator_count >= 4


def test_export_materials_to_docx_avoids_overwriting_existing_file(tmp_path):
    first_path = export_materials_to_docx(
        materials_data=sample_materials_data(),
        user_profile=sample_user_profile(),
        output_dir=tmp_path,
        filename="materials.docx",
    )

    second_path = export_materials_to_docx(
        materials_data=sample_materials_data(),
        user_profile=sample_user_profile(),
        output_dir=tmp_path,
        filename="materials.docx",
    )

    assert Path(first_path).exists()
    assert Path(second_path).exists()
    assert first_path != second_path
    assert Path(second_path).name == "materials_1.docx"


def test_export_materials_to_docx_missing_resume_data_raises_error(tmp_path):
    materials_data = sample_materials_data()
    del materials_data["resume"]

    with pytest.raises(ValueError):
        export_materials_to_docx(
            materials_data=materials_data,
            user_profile=sample_user_profile(),
            output_dir=tmp_path,
            filename="materials.docx",
        )


def test_export_materials_to_docx_missing_user_profile_data_raises_error(tmp_path):
    user_profile = sample_user_profile()
    del user_profile["email"]

    with pytest.raises(ValueError):
        export_materials_to_docx(
            materials_data=sample_materials_data(),
            user_profile=user_profile,
            output_dir=tmp_path,
            filename="materials.docx",
        )


def test_export_materials_to_docx_missing_strengths_data_raises_error(tmp_path):
    materials_data = sample_materials_data()
    del materials_data["strengths"]

    with pytest.raises(ValueError):
        export_materials_to_docx(
            materials_data=materials_data,
            user_profile=sample_user_profile(),
            output_dir=tmp_path,
            filename="materials.docx",
        )