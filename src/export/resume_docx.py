# src/export/resume_docx.py

"""DOCX export layer for generated application materials.

Architecture position:
    interface -> export

This module converts generated application materials into a formatted
.docx file. The exported document includes a resume page, a cover letter
page, and a strengths and weaknesses page.

The resume page uses the user's profile information for the header and
education section, and uses generated materials for the tailored skills,
projects, and experience sections.

The exported file uses:
    - Garamond font.
    - A centered name header.
    - A centered contact line.
    - Section headings for education, skills, projects, and experience.
    - Separator lines between major resume sections.
    - Separate pages for the cover letter and strengths/weaknesses.

Error contract:
    This module raises exceptions instead of returning status dictionaries
    because it is an export helper. The interface layer catches these
    exceptions and displays user-facing export messages.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, cast

from docx import Document as create_document
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


FONT_NAME = "Garamond"
SEPARATOR_WEIGHT_EIGHTHS_OF_POINT = "18"


def _set_run_font(
    run,
    size: int = 12,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """Apply the standard font formatting to a Word run.

    Parameters:
        run (docx.text.run.Run): Word run object to format.
        size (int): Font size in points. Defaults to 12.
        bold (bool): True if the run should be bold; otherwise, False.
        italic (bool): True if the run should be italic; otherwise, False.

    Returns:
        None: This function modifies the run in place.
    """

    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

    if run._element.rPr is not None and run._element.rPr.rFonts is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def _set_normal_style(document: Document) -> None:
    """Set the default document font style.

    Parameters:
        document (Document): Word document object to update.

    Returns:
        None: This function modifies the document style in place.
    """

    style = cast(Any, document.styles["Normal"])

    style.font.name = FONT_NAME
    style.font.size = Pt(12)

    if style._element.rPr is not None and style._element.rPr.rFonts is not None:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def _add_separator_line(document: Document) -> None:
    """Add a horizontal separator line with weight 2 1/4 pt.

    Parameters:
        document (Document): Word document object where the separator line
        will be added.

    Returns:
        None: This function adds a paragraph with a bottom border to the
        document.
    """

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(2)
    paragraph_format.space_after = Pt(6)

    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = paragraph_properties.find(qn("w:pBdr"))

    if borders is None:
        borders = OxmlElement("w:pBdr")
        paragraph_properties.append(borders)

    bottom_border = borders.find(qn("w:bottom"))

    if bottom_border is None:
        bottom_border = OxmlElement("w:bottom")
        borders.append(bottom_border)

    bottom_border.set(qn("w:val"), "single")
    bottom_border.set(qn("w:sz"), SEPARATOR_WEIGHT_EIGHTHS_OF_POINT)
    bottom_border.set(qn("w:space"), "1")
    bottom_border.set(qn("w:color"), "000000")


def _add_centered_line(
    document: Document,
    text: str,
    size: int,
    bold: bool = False,
) -> None:
    """Add one centered line of text to the document.

    Parameters:
        document (Document): Word document object where the line will be added.
        text (str): Text to display.
        size (int): Font size in points.
        bold (bool): True if the text should be bold; otherwise, False.

    Returns:
        None: This function adds a formatted paragraph to the document.
    """

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(text)
    _set_run_font(run, size=size, bold=bold)


def _add_section_heading(document: Document, text: str) -> None:
    """Add an uppercase section heading to the document.

    Parameters:
        document (Document): Word document object where the heading will be
        added.
        text (str): Heading text before uppercase formatting is applied.

    Returns:
        None: This function adds a bold section heading paragraph to the
        document.
    """

    paragraph = document.add_paragraph()
    run = paragraph.add_run(text.upper())
    _set_run_font(run, size=12, bold=True)


def _add_bullet(document: Document, text: str) -> None:
    """Add one bullet item to the document.

    Parameters:
        document (Document): Word document object where the bullet will be
        added.
        text (str): Bullet text to display.

    Returns:
        None: This function adds a formatted bullet paragraph to the document.
    """

    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    _set_run_font(run, size=12)


def _add_bullet_section(document: Document, title: str, items: list[str]) -> None:
    """Add a heading followed by bullet items.

    Parameters:
        document (Document): Word document object where the section will be
        added.
        title (str): Section heading text.
        items (list[str]): Bullet items to include in the section.

    Returns:
        None: This function adds the section heading and bullet list to the
        document.
    """

    _add_section_heading(document, title)

    if not items:
        _add_bullet(document, "None provided.")
        return

    for item in items:
        _add_bullet(document, item)


def _add_resume_page(
    document: Document,
    materials_data: dict,
    user_profile: dict,
) -> None:
    """Add the formatted resume page to the document.

    Parameters:
        document (Document): Word document object where the resume page will
        be added.
        materials_data (dict): Generated application materials.

            Expected format:
                {
                    "resume": {
                        "skills": list[str],
                        "projects": list[str],
                        "experience": list[str],
                    },
                    "cover_letter": str,
                    "strengths": list[str],
                    "weaknesses": list[str],
                }

        user_profile (dict): Saved user profile data.

            Expected format:
                {
                    "name": str,
                    "email": str,
                    "phone_number": str,
                    "university": str,
                    "degree": str,
                }

    Returns:
        None: This function adds the resume content to the document.
    """

    resume = materials_data.get("resume", {})

    name = user_profile.get("name", "Unknown User")
    email = user_profile.get("email", "")
    phone_number = user_profile.get("phone_number", "")
    university = user_profile.get("university", "")
    degree = user_profile.get("degree", "")

    contact_parts = []

    if email:
        contact_parts.append(email)

    if phone_number:
        contact_parts.append(phone_number)

    contact_line = " | ".join(contact_parts)

    _add_centered_line(document, name, size=28, bold=True)

    if contact_line:
        _add_centered_line(document, contact_line, size=13, bold=False)

    _add_separator_line(document)

    _add_section_heading(document, "Education")

    education_paragraph = document.add_paragraph()

    university_run = education_paragraph.add_run(university)
    _set_run_font(university_run, size=12, bold=True)

    if university and degree:
        separator_run = education_paragraph.add_run(" - ")
        _set_run_font(separator_run, size=12)

    degree_run = education_paragraph.add_run(degree)
    _set_run_font(degree_run, size=12, italic=True)

    _add_separator_line(document)

    _add_bullet_section(document, "Skills", resume.get("skills", []))
    _add_separator_line(document)

    _add_bullet_section(document, "Projects", resume.get("projects", []))
    _add_separator_line(document)

    _add_bullet_section(document, "Experience", resume.get("experience", []))


def _add_cover_letter_page(document: Document, cover_letter: str) -> None:
    """Add the cover letter page to the document.

    Parameters:
        document (Document): Word document object where the cover letter page
        will be added.
        cover_letter (str): Generated cover letter text.

    Returns:
        None: This function adds the cover letter heading and paragraphs to
        the document.
    """

    _add_section_heading(document, "Cover Letter")

    paragraphs = cover_letter.split("\n")

    for paragraph_text in paragraphs:
        paragraph_text = paragraph_text.strip()

        if not paragraph_text:
            document.add_paragraph()
            continue

        paragraph = document.add_paragraph()
        run = paragraph.add_run(paragraph_text)
        _set_run_font(run, size=12)


def _add_strengths_and_weaknesses_page(
    document: Document,
    strengths: list[str],
    weaknesses: list[str],
) -> None:
    """Add the strengths and weaknesses page to the document.

    Parameters:
        document (Document): Word document object where the page will be added.
        strengths (list[str]): Generated strengths for the target job.
        weaknesses (list[str]): Generated weaknesses and preparation advice for
        the target job.

    Returns:
        None: This function adds strengths and weaknesses sections to the
        document.
    """

    _add_section_heading(document, "Strengths")

    if not strengths:
        _add_bullet(document, "No strengths were generated.")
    else:
        for strength in strengths:
            _add_bullet(document, strength)

    _add_separator_line(document)

    _add_section_heading(document, "Weaknesses")

    if not weaknesses:
        _add_bullet(document, "No weaknesses were generated.")
    else:
        for weakness in weaknesses:
            _add_bullet(document, weakness)


def _get_available_file_path(output_dir: str | Path, filename: str) -> Path:
    """Return an available file path without overwriting existing files.

    Parameters:
        output_dir (str | Path): Directory where the file should be saved.
        filename (str): Requested file name.

    Returns:
        Path: Available path for the requested file. If the requested file
        already exists, the function returns a numbered version such as
        "materials_1.docx".
    """

    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    file_path = output_path / filename

    if not file_path.exists():
        return file_path

    stem = file_path.stem
    suffix = file_path.suffix

    counter = 1

    while True:
        candidate = output_path / f"{stem}_{counter}{suffix}"

        if not candidate.exists():
            return candidate

        counter += 1


def _validate_materials_data(materials_data: dict) -> None:
    """Validate generated materials before exporting to DOCX.

    Parameters:
        materials_data (dict): Generated application materials.

        Expected format:
            {
                "resume": {
                    "skills": list[str],
                    "projects": list[str],
                    "experience": list[str],
                },
                "cover_letter": str,
                "strengths": list[str],
                "weaknesses": list[str],
            }

    Returns:
        None: This function returns None when validation succeeds.

    Raises:
        ValueError: If generated materials are missing.
        ValueError: If resume data is missing.
        ValueError: If resume sections are incomplete.
        ValueError: If cover letter data is missing.
        ValueError: If strengths data is missing.
        ValueError: If weaknesses data is missing.
    """

    if not isinstance(materials_data, dict):
        raise ValueError("Generated materials are required.")

    resume = materials_data.get("resume")

    if not isinstance(resume, dict):
        raise ValueError("Resume data is required.")

    for key in ["skills", "projects", "experience"]:
        if key not in resume or not isinstance(resume[key], list):
            raise ValueError("Resume data is incomplete.")

    if not isinstance(materials_data.get("cover_letter"), str):
        raise ValueError("Cover letter is required.")

    if not isinstance(materials_data.get("strengths"), list):
        raise ValueError("Strengths data is required.")

    if not isinstance(materials_data.get("weaknesses"), list):
        raise ValueError("Weaknesses data is required.")


def _validate_user_profile(user_profile: dict) -> None:
    """Validate user profile data before exporting to DOCX.

    Parameters:
        user_profile (dict): Saved user profile data.

        Expected format:
            {
                "name": str,
                "email": str,
                "phone_number": str,
                "university": str,
                "degree": str,
            }

    Returns:
        None: This function returns None when validation succeeds.

    Raises:
        ValueError: If user profile data is missing.
        ValueError: If required user profile fields are incomplete.
    """

    if not isinstance(user_profile, dict):
        raise ValueError("User profile is required.")

    required_keys = {
        "name",
        "email",
        "phone_number",
        "university",
        "degree",
    }

    if not required_keys.issubset(user_profile.keys()):
        raise ValueError("User profile is incomplete.")


def export_materials_to_docx(
    materials_data: dict,
    user_profile: dict,
    output_dir: str | Path = "outputs",
    filename: str | None = None,
) -> str:
    """Export resume, cover letter, strengths, and weaknesses to one .docx file.

    Parameters:
        materials_data (dict): Generated application materials.

            Expected format:
                {
                    "resume": {
                        "skills": list[str],
                        "projects": list[str],
                        "experience": list[str],
                    },
                    "cover_letter": str,
                    "strengths": list[str],
                    "weaknesses": list[str],
                }

        user_profile (dict): Saved user profile data used for the resume
        header and education section.

            Expected format:
                {
                    "name": str,
                    "email": str,
                    "phone_number": str,
                    "university": str,
                    "degree": str,
                }

        output_dir (str | Path): Folder where the .docx file should be saved.
        Defaults to "outputs".

        filename (str | None): Optional file name. If None, the function uses
        "application_materials.docx".

    Returns:
        str: Path to the saved .docx file.

    Raises:
        ValueError: If materials_data is missing required fields.
        ValueError: If user_profile is missing required fields.
        PermissionError: If the file cannot be written because it is open or
        blocked by the operating system.
        Exception: If python-docx cannot create or save the file for another
        reason.
    """

    _validate_materials_data(materials_data)
    _validate_user_profile(user_profile)

    if filename is None:
        filename = "application_materials.docx"

    document = create_document()
    _set_normal_style(document)

    _add_resume_page(document, materials_data, user_profile)

    document.add_page_break()
    _add_cover_letter_page(document, materials_data["cover_letter"])

    document.add_page_break()
    _add_strengths_and_weaknesses_page(
        document,
        materials_data["strengths"],
        materials_data["weaknesses"],
    )

    file_path = _get_available_file_path(output_dir, filename)
    document.save(str(file_path))

    return str(file_path)